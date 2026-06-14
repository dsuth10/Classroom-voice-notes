import sys
import os
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def convert_docx_to_md(input_path, output_path):
    doc = Document(input_path)
    md_content = []

    # Iterate through all elements in the body
    for element in doc.element.body:
        if element.tag.endswith('p'): # Paragraph
            para = Paragraph(element, doc)
            text = get_para_markdown(para)
            if not text.strip():
                continue
            
            # Map styles to markdown
            style = para.style.name
            if 'Heading 1' in style:
                md_content.append(f"# {text}\n")
            elif 'Heading 2' in style:
                md_content.append(f"## {text}\n")
            elif 'Heading 3' in style:
                md_content.append(f"### {text}\n")
            elif 'List Bullet' in style or para._element.xpath('.//w:numPr'):
                # Handle nested lists if possible, otherwise simple bullet
                md_content.append(f"- {text}\n")
            else:
                md_content.append(f"{text}\n")
        
        elif element.tag.endswith('tbl'): # Table
            table = Table(element, doc)
            md_content.append(convert_table_to_md(table))

    # Clean up multiple blank lines
    final_output = "\n".join(md_content)
    final_output = re.sub(r'\n{3,}', '\n\n', final_output)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(final_output)

def get_para_markdown(para):
    """Converts a paragraph to Markdown, preserving bold and italic runs."""
    md_runs = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        
        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        md_runs.append(text)
    
    return "".join(md_runs)

def is_table_complex(table):
    """Detects if a table has merged cells or irregular structure."""
    try:
        first_row_cells = len(table.rows[0].cells)
        for row in table.rows:
            if len(row.cells) != first_row_cells:
                return True
            for cell in row.cells:
                # Check for merged cells (gridSpan or vMerge)
                tc = cell._tc
                grid_span = tc.xpath('./w:tcPr/w:gridSpan')
                v_merge = tc.xpath('./w:tcPr/w:vMerge')
                if grid_span or v_merge:
                    return True
    except:
        # If we can't even iterate rows/cells properly, it's definitely complex/irregular
        return True
    return False

def convert_table_to_md(table):
    if not is_table_complex(table):
        # Regular table -> Pipe table
        rows = []
        try:
            for i, row in enumerate(table.rows):
                # We use the paragraph parser for each cell to preserve bold/italic
                cell_texts = []
                for cell in row.cells:
                    cell_md = " ".join([get_para_markdown(p) for p in cell.paragraphs]).replace('\n', '<br>').strip()
                    cell_texts.append(cell_md)
                
                rows.append("| " + " | ".join(cell_texts) + " |")
                if i == 0:
                    rows.append("| " + " | ".join(["---"] * len(cell_texts)) + " |")
            return "\n" + "\n".join(rows) + "\n"
        except:
            # Fallback to complex handling if pipe table fails
            pass

    # Complex table -> Bulleted sections (the "Structure-Aware" fallback)
    lines = ["\n> [TABLE STRUCTURE]"]
    for i, row in enumerate(table.rows):
        lines.append(f"- **Row {i+1}:**")
        for j, cell in enumerate(row.cells):
            # Collect markdown text for the cell
            cell_md = " ".join([get_para_markdown(p) for p in cell.paragraphs]).strip()
            if cell_md:
                # If cell has multiple paragraphs, indent them
                cell_lines = cell_md.split('\n')
                lines.append(f"  - Column {j+1}: {cell_lines[0]}")
                for extra in cell_lines[1:]:
                    lines.append(f"    {extra}")
    return "\n" + "\n".join(lines) + "\n"

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python structure_aware_convert.py <input.docx> <output.md>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(os.path.dirname(output_file)):
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
    convert_docx_to_md(input_file, output_file)
